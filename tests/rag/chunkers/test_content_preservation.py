from docx import Document

from src.rag_platform.rag.chunkers.manual_chunker import ManualChunker
from src.rag_platform.rag.chunkers.rule_chunker import RuleChunker
from src.rag_platform.rag.chunkers.sop_chunker import SopChunker
from src.rag_platform.rag.parsers.rule_docx_parser import RuleDocxParser


def test_sop_parent_chunk_preserves_precheck_exception_and_upgrade_text() -> None:
    clean_content = (
        "仅退款申请处理SOP\n"
        "适用场景：订单未发货。\n"
        "前置检查：核对订单状态和用户凭证。\n"
        "1. 打开售后工单。\n"
        "2. 核对退款原因。\n"
        "异常分支：凭证缺失时暂停处理。\n"
        "升级条件：高金额订单升级人工复核。\n"
        "关键事实：退款前必须核对订单状态。"
    )
    chunks = SopChunker().build_chunks(
        structure={
            "title": "仅退款申请处理SOP",
            "scene": "适用场景：订单未发货。",
            "steps": ["1. 打开售后工单。", "2. 核对退款原因。"],
            "notes": [],
        },
        clean_content=clean_content,
    )

    parent = chunks[0]
    assert "前置检查：核对订单状态和用户凭证。" in parent.content
    assert "异常分支：凭证缺失时暂停处理。" in parent.content
    assert "升级条件：高金额订单升级人工复核。" in parent.content
    assert "关键事实：退款前必须核对订单状态。" in parent.content


def test_manual_chunk_preserves_menu_fields_errors_and_fact_text() -> None:
    clean_content = (
        "退款管理后台操作手册\n"
        "菜单路径：退款管理 > 退款单查询。\n"
        "字段说明：退款单号、支付流水号。\n"
        "1. 点击【查询】按钮。\n"
        "2. 点击【重试退款】按钮。\n"
        "错误提示：E-RM-8101表示退款单不存在。\n"
        "关键事实：可按退款单号查询执行状态。"
    )
    chunks = ManualChunker().build_chunks(
        structure={
            "title": "退款管理后台操作手册",
            "title_path": "退款管理 > 退款单查询",
            "steps": ["1. 点击【查询】按钮。", "2. 点击【重试退款】按钮。"],
            "button_names": ["查询", "重试退款"],
        },
        clean_content=clean_content,
    )

    assert len(chunks) == 1
    content = chunks[0].content
    assert "菜单路径：退款管理 > 退款单查询。" in content
    assert "字段说明：退款单号、支付流水号。" in content
    assert "错误提示：E-RM-8101表示退款单不存在。" in content
    assert "关键事实：可按退款单号查询执行状态。" in content


def test_rule_parser_preserves_multiline_clause_content(tmp_path) -> None:
    path = tmp_path / "rule.docx"
    document = Document()
    document.add_paragraph("优惠券使用规则 V1")
    document.add_paragraph(
        "2. 规则条款：每个订单仅可使用一张平台券。\n"
        "关键事实：V1每个订单仅可使用一张平台券"
    )
    document.save(path)

    result = RuleDocxParser().parse(str(path))

    assert (
        "关键事实：V1每个订单仅可使用一张平台券"
        in result.structure["clauses"][0]["content"]
    )


def test_rule_chunk_preserves_raw_multiline_clause_content() -> None:
    chunks = RuleChunker().build_chunks(
        structure={
            "title": "优惠券使用规则 V1",
            "clauses": [
                {
                    "clause_no": "2",
                    "title_path": "",
                    "content": "规则条款：每个订单仅可使用一张平台券。",
                    "raw_line": (
                        "2. 规则条款：每个订单仅可使用一张平台券。\n"
                        "关键事实：V1每个订单仅可使用一张平台券"
                    ),
                }
            ],
        },
        clean_content="",
    )

    assert "V1每个订单仅可使用一张平台券" in chunks[0].content
