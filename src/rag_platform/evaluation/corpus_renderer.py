import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

from src.rag_platform.domain.document import DocumentType
from src.rag_platform.evaluation.corpus_models import (
    DocumentBlueprint,
    GeneratedDocumentSection,
    GeneratedSourceDocument,
)
from src.rag_platform.rag.parsers.parser_factory import DocumentParserFactory


@dataclass
class RenderVerification:
    extracted_text: str
    errors: list[str] = field(default_factory=list)

    @property
    def is_valid(self) -> bool:
        return not self.errors


class CorpusRenderer:
    def __init__(self, pdf_font_path: str | None = None) -> None:
        self.pdf_font_name = self._register_pdf_font(pdf_font_path)

    def render(
        self,
        document: GeneratedSourceDocument,
        output_dir: Path,
    ) -> Path:
        output_dir.mkdir(parents=True, exist_ok=True)
        if document.doc_type in {
            DocumentType.FAQ,
            DocumentType.RULE,
        }:
            output_path = output_dir / f"{document.source_doc_code}.docx"
            self._render_docx(document, output_path)
            return output_path

        output_path = output_dir / f"{document.source_doc_code}.pdf"
        self._render_pdf(document, output_path)
        return output_path

    def verify(
        self,
        blueprint: DocumentBlueprint,
        document: GeneratedSourceDocument,
        output_path: Path,
    ) -> RenderVerification:
        parser = DocumentParserFactory().get_parser(
            DocumentType(document.doc_type.value)
        )
        result = parser.parse(str(output_path))
        extracted = result.raw_content
        normalized_extracted = _normalize(extracted)
        errors: list[str] = []

        for identifier in blueprint.required_identifiers:
            if _normalize(identifier) not in normalized_extracted:
                errors.append(f"渲染后缺少必要标识符 {identifier}")
        for fact_text in document.fact_texts():
            if _normalize(fact_text) not in normalized_extracted:
                errors.append(f"渲染后缺少事实文本 {fact_text}")

        if document.doc_type == DocumentType.FAQ:
            if len(result.structure.get("qa_pairs", [])) != len(document.sections):
                errors.append("FAQ 渲染后问答数量与源 JSON 不一致")
        elif document.doc_type in {DocumentType.SOP, DocumentType.MANUAL}:
            if not result.structure.get("steps"):
                errors.append(f"{document.doc_type.value} 渲染后未解析到步骤")
        elif document.doc_type == DocumentType.RULE:
            if not result.structure.get("clauses"):
                errors.append("RULE 渲染后未解析到条款")

        return RenderVerification(extracted_text=extracted, errors=errors)

    def _render_docx(
        self,
        document: GeneratedSourceDocument,
        output_path: Path,
    ) -> None:
        doc = Document()
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.start_type = WD_SECTION.NEW_PAGE

        self._configure_docx_styles(doc)
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(document.title)
        self._set_run_font(title_run, 20, bold=True, color="1F4D78")

        metadata = doc.add_paragraph()
        metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
        metadata_run = metadata.add_run(
            f"文档编码：{document.source_doc_code}  |  "
            f"版本：{document.version}  |  主题：{document.topic}"
        )
        self._set_run_font(metadata_run, 9, color="666666")

        if document.effective_from:
            effective = doc.add_paragraph(
                f"生效日期：{document.effective_from.isoformat()}"
            )
            effective.style = doc.styles["Normal"]

        for index, item in enumerate(document.sections, start=1):
            if document.doc_type == DocumentType.FAQ:
                self._add_faq_section(doc, item, index)
            else:
                self._add_rule_section(doc, item, index)
        doc.save(output_path)

    def _add_faq_section(
        self,
        doc: Document,
        section: GeneratedDocumentSection,
        index: int,
    ) -> None:
        question = doc.add_paragraph(style="Heading 2")
        question.add_run(f"Q{index}：{section.heading}")

        answer_text = _append_fact_texts(section.content, section)
        answer = doc.add_paragraph(style="Normal")
        answer.add_run(f"A：{answer_text}")

        aliases = doc.add_paragraph(style="Normal")
        aliases.add_run(f"同义问法：{'；'.join(section.aliases)}")

    def _add_rule_section(
        self,
        doc: Document,
        section: GeneratedDocumentSection,
        index: int,
    ) -> None:
        content = _append_fact_texts(section.content, section)
        paragraph = doc.add_paragraph(style="Normal")
        paragraph.add_run(f"{index}. {section.heading}：{content}")

    def _configure_docx_styles(self, doc: Document) -> None:
        normal = doc.styles["Normal"]
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.25
        normal.paragraph_format.keep_together = True
        self._set_style_font(normal, "PingFang SC")

        heading = doc.styles["Heading 2"]
        heading.font.size = Pt(13)
        heading.font.bold = True
        heading.font.color.rgb = RGBColor.from_string("2E74B5")
        heading.paragraph_format.space_before = Pt(14)
        heading.paragraph_format.space_after = Pt(7)
        self._set_style_font(heading, "PingFang SC")

    def _render_pdf(
        self,
        document: GeneratedSourceDocument,
        output_path: Path,
    ) -> None:
        pdf = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=inch,
            leftMargin=inch,
            topMargin=inch,
            bottomMargin=inch,
            title=document.title,
            author="rag-platform evaluation corpus",
        )
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "CorpusTitle",
            parent=styles["Title"],
            fontName=self.pdf_font_name,
            fontSize=20,
            leading=26,
            textColor="#1F4D78",
            spaceAfter=12,
        )
        metadata_style = ParagraphStyle(
            "CorpusMetadata",
            parent=styles["Normal"],
            fontName=self.pdf_font_name,
            fontSize=9,
            leading=14,
            textColor="#666666",
            alignment=TA_LEFT,
            spaceAfter=12,
        )
        heading_style = ParagraphStyle(
            "CorpusHeading",
            parent=styles["Heading2"],
            fontName=self.pdf_font_name,
            fontSize=13,
            leading=18,
            textColor="#2E74B5",
            spaceBefore=10,
            spaceAfter=6,
            keepWithNext=True,
        )
        body_style = ParagraphStyle(
            "CorpusBody",
            parent=styles["BodyText"],
            fontName=self.pdf_font_name,
            fontSize=10.5,
            leading=16,
            spaceAfter=5,
            wordWrap="CJK",
        )

        story = [
            Paragraph(_escape(document.title), title_style),
            Paragraph(
                _escape(
                    f"文档编码：{document.source_doc_code} | "
                    f"版本：{document.version} | 主题：{document.topic}"
                ),
                metadata_style,
            ),
        ]
        for section in document.sections:
            story.append(Paragraph(_escape(section.heading), heading_style))
            content = _append_fact_texts(section.content, section)
            for line in content.splitlines():
                if line.strip():
                    story.append(Paragraph(_escape(line.strip()), body_style))
            story.append(Spacer(1, 4))
        pdf.build(story)

    def _register_pdf_font(self, pdf_font_path: str | None) -> str:
        if pdf_font_path:
            path = Path(pdf_font_path)
            if not path.is_file():
                raise ValueError(f"EVAL_PDF_FONT_PATH 不存在：{path}")
            font_name = "EvalCJK"
            pdfmetrics.registerFont(TTFont(font_name, str(path)))
            return font_name

        font_name = "STSong-Light"
        if font_name not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(UnicodeCIDFont(font_name))
        return font_name

    @staticmethod
    def _set_style_font(style, font_name: str) -> None:
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:ascii"), font_name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    @staticmethod
    def _set_run_font(
        run,
        size: int,
        *,
        bold: bool = False,
        color: str | None = None,
    ) -> None:
        run.font.name = "PingFang SC"
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "PingFang SC")
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "PingFang SC")
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "PingFang SC")
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def _append_fact_texts(
    content: str,
    section: GeneratedDocumentSection,
) -> str:
    result = content.strip()
    normalized = _normalize(result)
    missing = [
        fact.fact_text
        for fact in section.facts
        if _normalize(fact.fact_text) not in normalized
    ]
    if missing:
        result += "\n关键事实：" + "；".join(missing)
    return result


def _normalize(value: str) -> str:
    return re.sub(r"\s+", "", value).casefold()


def _escape(value: str) -> str:
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
