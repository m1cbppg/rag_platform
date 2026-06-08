import re

from docx import Document

from src.rag_platform.rag.parsers.base import BaseDocumentParser, ParseResult


class FaqDocxParser(BaseDocumentParser):
    """
    FAQ docx 解析器。

    支持两种 FAQ 格式：

    1. 表格型：
       问题 | 答案 | 同义问法 | 标签

    2. 段落型：
       Q1：车辆定位不刷新怎么办？
       A：客服应先确认用户是否绑定车辆...
       同义问法：车辆定位不刷新怎么处理？；车主反馈车辆定位不刷新怎么办
    """

    # 匹配问题行：
    # Q：xxx
    # Q1：xxx
    # Q10: xxx
    # 问题：xxx
    # 问题1：xxx
    QUESTION_PATTERN = re.compile(r"^(Q\d*|问题\d*)[:：]\s*(.+)$", re.IGNORECASE)

    # 匹配答案行：
    # A：xxx
    # A1：xxx
    # 答案：xxx
    ANSWER_PATTERN = re.compile(r"^(A\d*|答案\d*)[:：]\s*(.+)$", re.IGNORECASE)

    # 匹配同义问法行：
    # 同义问法：xxx；xxx；xxx
    ALIAS_PATTERN = re.compile(r"^(同义问法|相似问法|别名|aliases?)[:：]\s*(.+)$", re.IGNORECASE)

    def parse(self, file_path: str) -> ParseResult:
        """
        解析 FAQ docx 文件。

        重点：
        raw_content 保存原始 docx 文本；
        structure 保存解析出来的 qa_pairs 和 metadata。
        """

        doc = Document(file_path)

        raw_content = self._extract_raw_content(doc)
        metadata = self._parse_metadata_tables(doc)

        qa_pairs = self._parse_tables(doc)

        if not qa_pairs:
            qa_pairs = self._parse_paragraphs(doc)

        return ParseResult(
            raw_content=raw_content,
            structure={
                "doc_type": "FAQ",
                "metadata": metadata,
                "qa_pairs": qa_pairs,
            },
            parser_type="FAQ_DOCX",
        )

    def _extract_raw_content(self, doc: Document) -> str:
        """
        提取 docx 原始文本。

        为什么不用 qa_pairs 拼 raw_content？
        因为 qa_pairs 解析失败时，raw_content 也会变空。
        raw_content 应该尽量保留原始可见文本，方便排查。
        """

        parts: list[str] = []

        # 提取段落文本
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)

    def _parse_metadata_tables(self, doc: Document) -> dict:
        """
        解析文档元数据表格。

        你的文件里有类似：
        文档版本 | 2026.06
        业务域 | 车辆定位/设备
        推荐切分 | Q/A 段落为一个 chunk
        典型召回词 | 定位不刷新、设备离线、GPS 漂移、SIM 卡欠费

        这类表格不是 FAQ 问答表，而是元数据表。
        """

        metadata: dict[str, str] = {}

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]

                if len(cells) >= 2:
                    key = cells[0]
                    value = cells[1]

                    if key and value:
                        metadata[key] = value

        return metadata

    def _parse_tables(self, doc: Document) -> list[dict]:
        """
        从 docx 表格中解析 FAQ。

        当前文件不是 FAQ 表格型，所以这里一般不会解析出 qa_pairs。
        但是保留这个能力，方便以后支持表格 FAQ。
        """

        qa_pairs: list[dict] = []

        for table in doc.tables:
            rows = table.rows

            if not rows:
                continue

            headers = [cell.text.strip() for cell in rows[0].cells]

            question_idx = self._find_col_index(headers, ["问题", "question", "Q"])
            answer_idx = self._find_col_index(headers, ["答案", "answer", "A"])
            aliases_idx = self._find_col_index(headers, ["同义问法", "相似问法", "aliases"])
            tags_idx = self._find_col_index(headers, ["标签", "tags"])

            if question_idx is None or answer_idx is None:
                continue

            for row in rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]

                question = self._safe_get(cells, question_idx)
                answer = self._safe_get(cells, answer_idx)

                if not question or not answer:
                    continue

                aliases = self._split_multi_value(self._safe_get(cells, aliases_idx))
                tags = self._split_multi_value(self._safe_get(cells, tags_idx))

                qa_pairs.append({
                    "question": question,
                    "answer": answer,
                    "aliases": aliases,
                    "tags": tags,
                })

        return qa_pairs

    def _parse_paragraphs(self, doc: Document) -> list[dict]:
        """
        从段落中解析 FAQ。

        支持：
        Q1：问题
        A：答案
        同义问法：xxx；xxx
        """

        qa_pairs: list[dict] = []

        current_question: str | None = None
        current_answer_parts: list[str] = []
        current_aliases: list[str] = []

        for paragraph in doc.paragraphs:
            line = paragraph.text.strip()

            if not line:
                continue

            question_match = self.QUESTION_PATTERN.match(line)
            answer_match = self.ANSWER_PATTERN.match(line)
            alias_match = self.ALIAS_PATTERN.match(line)

            if question_match:
                # 遇到新问题时，先保存上一组 Q/A
                if current_question and current_answer_parts:
                    qa_pairs.append({
                        "question": current_question,
                        "answer": "\n".join(current_answer_parts).strip(),
                        "aliases": current_aliases,
                        "tags": [],
                    })

                current_question = question_match.group(2).strip()
                current_answer_parts = []
                current_aliases = []

            elif answer_match:
                current_answer_parts.append(answer_match.group(2).strip())

            elif alias_match:
                current_aliases = self._split_multi_value(alias_match.group(2).strip())

            else:
                # 如果已经进入某个问题，普通段落默认追加到答案中
                if current_question:
                    current_answer_parts.append(line)

        # 保存最后一组 Q/A
        if current_question and current_answer_parts:
            qa_pairs.append({
                "question": current_question,
                "answer": "\n".join(current_answer_parts).strip(),
                "aliases": current_aliases,
                "tags": [],
            })

        return qa_pairs

    def _find_col_index(self, headers: list[str], candidates: list[str]) -> int | None:
        """
        根据候选列名查找表格列索引。
        """

        normalized_headers = [h.lower() for h in headers]

        for candidate in candidates:
            candidate_lower = candidate.lower()
            if candidate_lower in normalized_headers:
                return normalized_headers.index(candidate_lower)

        return None

    def _safe_get(self, values: list[str], index: int | None) -> str:
        """
        安全读取列表里的某个位置。
        """

        if index is None:
            return ""

        if index >= len(values):
            return ""

        return values[index].strip()

    def _split_multi_value(self, text: str) -> list[str]:
        """
        拆分多个同义问法或标签。

        支持：
        a；b；c
        a,b,c
        a，b，c
        a|b|c
        """

        if not text:
            return []

        for sep in ["，", "；", ";", "|"]:
            text = text.replace(sep, ",")

        return [item.strip() for item in text.split(",") if item.strip()]